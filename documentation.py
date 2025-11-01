from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

def create_project_documentation_docx(filename="CortexX_Demand_Forecasts_Report.docx"):
    """
    Creates comprehensive project documentation for CortexX Demand Forecasting Platform
    """
    doc = Document()
    
    # Set document properties
    doc_properties = doc.core_properties
    doc_properties.title = "CortexX Demand Forecasts Documentation"
    doc_properties.subject = "Project Documentation and Technical Specifications"
    doc_properties.author = "CortexX Development Team"
    doc_properties.keywords = "ML, Forecasting, Demand Prediction, Documentation"
    
    # Cover Page
    _create_cover_page(doc)
    
    # Table of Contents
    _create_table_of_contents(doc)
    
    # Main Content Sections
    _create_project_overview(doc)
    _create_technical_architecture(doc)
    _create_stakeholder_analysis(doc)
    _create_data_flow_design(doc)
    _create_ui_ux_design(doc)
    _create_implementation_plan(doc)
    
    # Save document
    doc.save(filename)
    print(f"✅ Professional project documentation saved as {filename}")

def _create_cover_page(doc):
    """Creates a professional cover page"""
    # Title Section
    title_section = doc.add_section()
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Main Title
    title_run = title_paragraph.add_run("CortexX Demand Forecasts\n")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = "Calibri"
    
    # Subtitle
    subtitle_run = title_paragraph.add_run("Technical Documentation & Project Specifications\n\n")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.name = "Calibri"
    
    # Version and Date
    version_run = title_paragraph.add_run(f"Version 1.0\n{datetime.now().strftime('%B %d, %Y')}")
    version_run.font.size = Pt(12)
    version_run.italic = True
    version_run.font.name = "Calibri"
    
    doc.add_page_break()

def _create_table_of_contents(doc):
    """Creates a table of contents"""
    doc.add_heading("Table of Contents", level=1)
    
    sections = [
        ("1. Project Overview", 2),
        ("   1.1 Purpose & Vision", 3),
        ("   1.2 Project Objectives", 3),
        ("   1.3 Scope & Deliverables", 3),
        ("2. Technical Architecture", 2),
        ("   2.1 System Overview", 3),
        ("   2.2 Technology Stack", 3),
        ("3. Stakeholder Analysis", 2),
        ("4. Data Flow Design", 2),
        ("5. UI/UX Design", 2),
        ("   5.1 Design Concept", 3),
        ("   5.2 User Flow", 3),
        ("   5.3 Wireframes", 3),
        ("6. Implementation Plan", 2)
    ]
    
    for section, level in sections:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2 * (level - 2))
        run = paragraph.add_run(section)
        if level == 2:
            run.bold = True
    
    doc.add_page_break()

def _create_project_overview(doc):
    """Creates the project overview section"""
    doc.add_heading("1. Project Overview", level=1)
    
    # Purpose & Vision
    doc.add_heading("1.1 Purpose & Vision", level=2)
    purpose_text = (
        "CortexX is an enterprise-level sales forecasting and demand prediction platform designed to "
        "help businesses accurately predict sales, analyze trends, and optimize inventory and staffing "
        "decisions using advanced machine learning techniques. The platform transforms raw sales data "
        "into actionable insights through automated pipelines and interactive visualization."
    )
    doc.add_paragraph(purpose_text)
    
    # Objectives
    doc.add_heading("1.2 Project Objectives", level=2)
    objectives = [
        "Build a scalable, modular forecasting tool adaptable to various business domains",
        "Offer multiple ML models (XGBoost, LightGBM, Prophet, ensembles) for flexible forecasting",
        "Provide comprehensive exploratory data analysis and automated feature engineering",
        "Deliver an interactive Streamlit dashboard for real-time insights and reporting",
        "Support automated model training, retraining, and performance monitoring",
        "Ensure maintainability through clean code architecture and comprehensive documentation"
    ]
    for objective in objectives:
        paragraph = doc.add_paragraph(objective, style='List Bullet')
        paragraph.paragraph_format.space_after = Pt(6)
    
    # Scope & Deliverables
    doc.add_heading("1.3 Scope & Deliverables", level=2)
    
    doc.add_heading("Functional Scope", level=3)
    scope_items = [
        "Data ingestion and pre-processing of sales and inventory data from CSV/Excel sources",
        "Automated feature creation for time-series forecasting including lag features and rolling statistics",
        "Training, evaluation, and comparison of sophisticated ML models with hyperparameter tuning",
        "Interactive visualization and reporting with export capabilities",
        "Flexible date handling and data validation for diverse business scenarios"
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Key Deliverables", level=3)
    deliverables = [
        "Installation and setup scripts for seamless deployment",
        "Well-documented data pipeline and model training modules",
        "Real-time dashboard with data upload, analysis, and forecasting capabilities",
        "Exportable forecast reports and model performance documentation",
        "Comprehensive API documentation and user guides"
    ]
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')

def _create_technical_architecture(doc):
    """Creates the technical architecture section"""
    doc.add_heading("2. Technical Architecture", level=1)
    
    # System Overview
    doc.add_heading("2.1 System Overview", level=2)
    system_overview = (
        "CortexX follows a modular microservices architecture with clear separation of concerns. "
        "The system is organized into distinct modules for data processing, feature engineering, "
        "model training, and visualization, enabling independent development and testing."
    )
    doc.add_paragraph(system_overview)
    
    # Technology Stack
    doc.add_heading("2.2 Technology Stack", level=2)
    tech_stack = [
        ("Backend & ML", "Python, Scikit-learn, XGBoost, LightGBM, Prophet, Pandas, NumPy"),
        ("Visualization", "Streamlit, Plotly, Matplotlib, Seaborn"),
        ("Data Processing", "Pandas, NumPy, SciPy for statistical operations"),
        ("Development", "Git, Poetry for dependency management, Pytest for testing"),
        ("Deployment", "Docker, Streamlit Cloud compatible")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "Technologies"
    
    for component, technologies in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = technologies

def _create_stakeholder_analysis(doc):
    """Creates the stakeholder analysis section"""
    doc.add_heading("3. Stakeholder Analysis", level=1)
    
    intro_text = (
        "Effective stakeholder management is crucial for project success. The following table "
        "outlines key stakeholders, their roles, and communication strategies."
    )
    doc.add_paragraph(intro_text)
    
    stakeholders = [
        ("Project Manager", "Oversees project delivery and timeline", "Weekly status meetings, email updates, milestone reviews"),
        ("Data Scientist", "Develops ML models and feature engineering", "Daily stand-ups, JIRA tickets, Slack coordination"),
        ("Software Developer", "Implements dashboard and backend services", "Code reviews, sprint demos, Slack channels"),
        ("Business Analyst", "Defines business requirements and success metrics", "Requirement workshops, status reports, email communication"),
        ("End Users", "Utilize forecasting platform for business decisions", "Training sessions, feedback collection, user support channels")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Stakeholder"
    hdr_cells[1].text = "Role & Responsibilities"
    hdr_cells[2].text = "Communication Plan"
    
    for name, role, comm in stakeholders:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = role
        row_cells[2].text = comm
    
    # Communication Note
    note_paragraph = doc.add_paragraph()
    note_paragraph.paragraph_format.space_before = Pt(12)
    note_run = note_paragraph.add_run("Communication Strategy: ")
    note_run.bold = True
    note_paragraph.add_run(
        "Scheduled meetings for formal updates, Slack channels for daily coordination, "
        "and email for official documentation and decision tracking."
    )

def _create_data_flow_design(doc):
    """Creates the data flow design section"""
    doc.add_heading("4. Data Flow Design", level=1)
    
    data_flow_text = (
        "CortexX employs a flexible file-based data ingestion system designed for rapid prototyping "
        "and deployment without complex database infrastructure requirements."
    )
    doc.add_paragraph(data_flow_text)
    
    doc.add_heading("Data Sources & Structure", level=2)
    data_sources = [
        "Primary data source: CSV files (e.g., retail_store_inventory.csv)",
        "Supported formats: CSV, Excel with flexible schema adaptation",
        "Key columns: Date, Store ID, Product ID, Category, Region, Inventory Level, Units Sold, Demand Forecast, Price, Discounts"
    ]
    for source in data_sources:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_heading("Data Processing Pipeline", level=2)
    pipeline_steps = [
        "Data Ingestion: Flexible CSV/Excel parsing with automatic date detection",
        "Pre-processing: Handling missing values, outliers, and data validation",
        "Feature Engineering: Automated creation of time-series features (lags, rolling statistics, seasonal patterns)",
        "Model Training: Multiple algorithm support with automated hyperparameter tuning",
        "Inference & Reporting: Forecast generation with confidence intervals and performance metrics"
    ]
    for step in pipeline_steps:
        doc.add_paragraph(step, style='List Number')

def _create_ui_ux_design(doc):
    """Creates the UI/UX design section"""
    doc.add_heading("5. UI/UX Design", level=1)
    
    # Design Concept
    doc.add_heading("5.1 Design Concept", level=2)
    design_principles = [
        "Intuitive navigation through sidebar-based workflow progression",
        "Consistent visual language using Plotly with clean, business-appropriate themes",
        "Progressive disclosure of complexity - simple defaults with advanced options",
        "Real-time feedback and validation throughout user interactions"
    ]
    for principle in design_principles:
        doc.add_paragraph(principle, style='List Bullet')
    
    # User Flow
    doc.add_heading("5.2 User Flow", level=2)
    user_flow_steps = [
        "Data Upload: Upload CSV/Excel or generate sample data for testing",
        "Configuration: Auto-detection or manual selection of date and value columns",
        "Exploration: Perform EDA with interactive visualizations and statistical summaries",
        "Feature Engineering: Automated feature creation with customization options",
        "Model Training: Configure and train forecasting models with selectable algorithms",
        "Forecasting: Generate and visualize sales forecasts with confidence intervals",
        "Reporting: Review model performance metrics and export comprehensive reports"
    ]
    for step in user_flow_steps:
        doc.add_paragraph(step, style='List Number')
    
    # Wireframes
    doc.add_heading("5.3 Wireframes Overview", level=2)
    wireframe_features = [
        "Grouped navigation tabs for logical workflow progression",
        "Dynamic sidebar displaying data state and session information",
        "Interactive data preview tables with sorting and filtering capabilities",
        "Multiple Plotly graph types supporting detailed exploratory analysis",
        "Model configuration panels with preset and custom options",
        "Export functionality for forecasts, charts, and performance reports"
    ]
    for feature in wireframe_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Design Rationale
    doc.add_heading("Design Rationale", level=3)
    rationale_points = [
        "Streamlit enables rapid development cycles and delivers intuitive user interfaces without complex frontend development",
        "Plotly provides enterprise-grade interactive visualization suitable for business intelligence applications",
        "Modular codebase supports maintainability, extensibility, and collaborative development practices",
        "Responsive design ensures accessibility across different devices and screen sizes"
    ]
    for point in rationale_points:
        doc.add_paragraph(point, style='List Bullet')

def _create_implementation_plan(doc):
    """Creates the implementation plan section"""
    doc.add_heading("6. Implementation Plan", level=1)
    
    doc.add_heading("Development Phases", level=2)
    phases = [
        ("Phase 1: Core Infrastructure", "Data pipeline, basic ML models, foundational dashboard"),
        ("Phase 2: Advanced Features", "Ensemble models, advanced EDA, performance optimization"),
        ("Phase 3: Production Ready", "Error handling, comprehensive testing, documentation"),
        ("Phase 4: Deployment & Training", "User training, deployment scripts, support materials")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Phase"
    hdr_cells[1].text = "Deliverables"
    
    for phase, deliverables in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = deliverables
    
    # Success Metrics
    doc.add_heading("Success Metrics", level=2)
    metrics = [
        "Model Accuracy: MAPE < 15% on validation datasets",
        "Performance: Dashboard load time < 3 seconds for standard datasets",
        "Usability: User satisfaction score > 4.0/5.0",
        "Reliability: System uptime > 99.5% in production environment"
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')

if __name__ == "__main__":
    create_project_documentation_docx()